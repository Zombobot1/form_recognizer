from skimage import measure
import cv2
import os
import numpy as np
import matplotlib.pyplot as plt
import matplotlib
from skimage.morphology import skeletonize
from skimage.color import rgb2gray
from PIL import Image
from pathlib import Path
from pdf2image import convert_from_path
from PIL import Image, ImageOps
from io import BytesIO
from docx import Document
import pytesseract
import io
import copy as cpy
from preproc import preprocess
# pwd = os.getcwd()+"/back"
# os.chdir(pwd)


def decodeImage(data):
    # Gives us 1d array
    decoded = np.fromstring(data, dtype=np.uint8)
    # We have to convert it into (270, 480,3) in order to see as an image
    # decoded = decoded.reshape((decoded.shape))
    decoded = rgb2gray(cv2.cvtColor(decoded, cv2.COLOR_BGR2RGB))
    return decoded


def recognize_cells(cells):
    results = []
    for row in cells:
        results.append([])
        column = 0
        for pic in row:
            column += 1

            copy = cpy.deepcopy(pic)
            img = Image.fromarray(pic)
            txt = None
            if column < 3:
                txt = pytesseract.image_to_string(
                    preprocess(copy), lang=os.getenv('MODEL'), config='-c tessedit_char_whitelist=" ,.0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyzÄÖÜßäöü" --psm 7')
            else:
                txt = pytesseract.image_to_string(
                    preprocess(copy), lang=os.getenv('MODEL'), config='-c tessedit_char_whitelist=0123456789,. --psm 8')

            txt = txt.replace('\n', ' ')
            results[-1].append((img, txt))
    return results


def save_to_docx(recognized_cells):
    return recognized_images_to_docx(recognized_cells)


def recognized_images_to_docx(recognized_images):
    def img_to_io(img) -> BytesIO:
        with BytesIO() as output:
            img.save(output, format="JPEG")
            r = BytesIO(output.getvalue())
        return r

    document = Document()

    max_i = len(recognized_images)
    max_j = len(recognized_images[0])

    table = document.add_table(rows=max_i, cols=max_j)
    table.style = 'Table Grid'
    print("started looping ")
    for i in range(max_i):
        for j in range(max_j):
            size = (300, 50)
            if j == 1 or j == 2:
                size = (1, 1)
            if j == 3:
                size = (100, 20)
            _img, text = recognized_images[i][j]
            _img.thumbnail(size, Image.LANCZOS)

            _img, text = recognized_images[i][j]
            img = img_to_io(_img)

            cell = table.rows[i].cells[j]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(img)
            cell.add_paragraph(text)

    result = BytesIO()
    document.save(result)
    result.seek(0)

    print("here is the doc ", document)
    return result


def write_bytesio_to_file(filename, bytesio):
    """
    Write the contents of the given BytesIO to a file.
    Creates the file or overwrites the file if it does
    not exist yet.
    """
    with open(filename, "wb") as outfile:
        # Copy the BytesIO stream to the output file
        outfile.write(bytesio.getbuffer())


# cells = pdf_to_cells('Bsp 1-5_geschwärzt-2-1.png')
# recognized_cells = recognize_cells(cells)
# print("starting saving ")

# file = save_to_docx(recognized_cells)
# write_bytesio_to_file('test.docx', file)

def find_table_head(img):
    cwd = os.getcwd()
    template_file = os.path.join(cwd + r'/template.png')
    template = cv2.imread(template_file, 0)  # template.shape is (102, 1456)
    th, tw = template.shape
    h, w = img.shape

    # Resize template if img scale is different than in the img the template was
    # taken from (with shape (2339, 1654)).
    if np.abs(h - 2339) > 2 or np.abs(w - 1654) > 2:
        new_th = int((102 * h) / 2339)
        new_tw = int((1456 * w) / 1654)
        template = cv2.resize(template, (new_tw, new_th))

    # Perform template matching without caring about scale invariance.
    result = cv2.matchTemplate(img, template, cv2.TM_CCOEFF_NORMED)
    (minVal, maxVal, minLoc, maxLoc) = cv2.minMaxLoc(result)
    (startX, startY) = maxLoc
    endX = startX + template.shape[1]
    endY = startY + template.shape[0]

    # Returns positions of the table head (startY is irrelevant).
    return startX, endX, endY


def extract_boxes(file):
    """
    Extracts the table cells of input image into a 2d array of image patches.
    Args:
        file (string):  Filepath to image file or pdf (only first page is
                        considered).
    Returns:
        result (numpy array):   Array of eqaul shape as input table, containing
                                table cells (numpy arrays). First and last row
                                are ignored.
    """
    # if Path(file).suffix == '.pdf':
    #     images = convert_from_path(file)
    #     img = np.array(images[0])
    #     img = (rgb2gray(img)*255).astype('uint8')
    # else:
    #     img = cv2.imread(file, 0)
    file_bytes = np.asarray(bytearray(file), dtype=np.uint8)

    img = cv2.imdecode(file_bytes, cv2.IMREAD_GRAYSCALE)
    h, w = img.shape

    # Binarize and invert.
    thresh, img_bin = cv2.threshold(img, 230, 255, cv2.THRESH_BINARY)
    #img = img_bin
    img_bin = 255-img_bin
    x0, x1, y0 = find_table_head(img_bin)

    # Remove upper and lower part of image + left and right margin.
    y0 += 3  # Make sure lower edge of table head is not included.
    img = img[y0:int(h*0.85), x0:x1]
    img_bin = img_bin[y0:int(h*0.85), x0:x1]

    # Detect vertical lines of at least 1/60 of total height.
    v_kernel_len = np.array(img).shape[1]//60
    # Detect horizontal lines of at least 1/20 of total width.
    h_kernel_len = np.array(img).shape[0]//20
    # Define a vertical kernel to detect all vertical lines of image.
    ver_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, v_kernel_len))
    # Define a horizontal kernel to detect all horizontal lines of image.
    hor_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (h_kernel_len, 1))

    # Use vertical kernel to detect the vertical lines.
    image_1 = cv2.erode(img_bin, ver_kernel, iterations=3)
    vertical_lines = cv2.dilate(image_1, ver_kernel, iterations=3)

    # Use horizontal kernel to detect the horizontal lines.
    image_2 = cv2.erode(img_bin, hor_kernel, iterations=3)
    horizontal_lines = cv2.dilate(image_2, hor_kernel, iterations=3)

    # Combine horizontal and vertical lines in a new third image, with both having same weight.
    img_vh = cv2.addWeighted(vertical_lines, 0.5, horizontal_lines, 0.5, 0.0)
    # Theshold and skeletonize the image.
    thresh, img_vh = cv2.threshold(
        img_vh, 128, 1, cv2.THRESH_BINARY | cv2.THRESH_OTSU)
    img_vh = skeletonize(img_vh)

    # Remove table lines from output image.
    # img[img_vh] = np.max(img)
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
    # img = cv2.dilate(img, kernel, iterations=1)
    # img = cv2.erode(img, kernel, iterations=1) # Close resulting gaps.
    # #
    # Remove table lines from binary image.
    img_bin[img_vh] = 0
    img_bin = cv2.erode(img_bin, kernel, iterations=1)
    # Close resulting gaps.
    img_bin = cv2.dilate(img_bin, kernel, iterations=2)
    img_bin = cv2.erode(img_bin, kernel, iterations=1)

    # Extract row indices.
    single_col = img_vh[:, img_vh.shape[1]//6]
    row_indices = np.nonzero(single_col)[0]
    row_indices = np.hstack((0, row_indices))

    # Calculate 1/3 of row height for detection of descenders of characters.
    row_height = row_indices[3] - row_indices[2]
    descenders_height = int(row_height * 0.33)

    # Extract column indices from middle of arbitrary row.
    single_row = img_vh[(row_indices[3]+row_indices[4])//2]
    col_indices = np.nonzero(single_row)[0]
    col_indices = np.hstack((0, col_indices, img_vh.shape[1]-1))

    n_rows, n_cols = 20, 4
    result = np.empty((n_rows, n_cols), dtype=object)

    for i in range(n_rows):
        for j in range(n_cols):
            if j == 0:
                # Label connected components of binary image patch.
                y0, y1 = row_indices[i]+1, row_indices[i+1] + descenders_height
                x0, x1 = col_indices[j]+1, col_indices[j+1]
                bin_patch = img_bin[y0:y1, x0:x1]
                labels = measure.label(bin_patch, background=0)
                # Check which components go through the lower case letter
                # compartment of patch.
                center_labels = np.unique(
                    labels[int(0.38*row_height):int(0.75*row_height)])
                center_labels = center_labels[center_labels != 0]
                # Don't change cell boundaries for empty cells
                if len(center_labels) > 2:
                    row_positions_letters = np.argwhere(
                        np.isin(labels, center_labels))[:, 0]
                    upper_bound = np.min(row_positions_letters)
                    lower_bound = np.max(row_positions_letters)
                else:
                    upper_bound = 0
                    lower_bound = y1 - y0
                # Cut out parts belonging to the cell above/below.
                img_patch = img[y0+upper_bound:y0+lower_bound, x0:x1]

            # No need to do this for numerical columns.
            else:
                img_patch = img[row_indices[i]+1:row_indices[i+1] +
                                2, col_indices[j]+1:col_indices[j+1]]

            result[i, j] = img_patch

    return result
